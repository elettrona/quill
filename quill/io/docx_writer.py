"""Native Word (.docx) writer built on python-docx.

The legacy export path runs Markdown through Pandoc, whose docx writer silently
drops the per-run font family, point size, color, highlight and per-paragraph
alignment that QUILL's hidden-codes feature
(``docs/rich-text-formatting-hidden-codes-design.md``) materializes. This module
consumes the shared :class:`~quill.io.rtf_model.RichDocument` and writes those
attributes onto real Word runs and paragraphs, so a styled QUILL document opens in
Word with its formatting intact and its structure screen-reader navigable.

python-docx is an optional dependency. :func:`python_docx_available` reports
whether it is importable; callers (``quill.io.export.write_docx_document``) fall
back to the Pandoc path when it is not, so docx export never hard-fails.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any

from quill.io.rtf_model import RichDocument, RichParagraph, markdown_to_rich

if TYPE_CHECKING:
    from quill.core.document import Document

__all__ = ["python_docx_available", "write_docx", "rich_to_docx_bytes"]

_NAMED_COLORS: dict[str, tuple[int, int, int]] = {
    "red": (255, 0, 0),
    "green": (0, 128, 0),
    "blue": (0, 0, 255),
    "black": (0, 0, 0),
    "white": (255, 255, 255),
    "yellow": (255, 255, 0),
    "orange": (255, 165, 0),
    "purple": (128, 0, 128),
    "gray": (128, 128, 128),
    "grey": (128, 128, 128),
}

#: python-docx highlight is an enum of named indices, not arbitrary RGB. Map the
#: common highlight names; anything else falls back to yellow.
_HIGHLIGHT_NAMES = {
    "yellow": "YELLOW",
    "green": "BRIGHT_GREEN",
    "bright_green": "BRIGHT_GREEN",
    "turquoise": "TURQUOISE",
    "pink": "PINK",
    "blue": "BLUE",
    "red": "RED",
    "gray": "GRAY_25",
    "grey": "GRAY_25",
}


def python_docx_available() -> bool:
    """Return ``True`` when python-docx can be imported."""
    try:
        import docx  # noqa: F401
    except ImportError:
        return False
    return True


def _parse_rgb(value: str) -> tuple[int, int, int] | None:
    text = value.strip()
    if text.startswith("#"):
        digits = text[1:]
        if len(digits) == 3:
            digits = "".join(char * 2 for char in digits)
        if len(digits) == 6:
            try:
                return (int(digits[0:2], 16), int(digits[2:4], 16), int(digits[4:6], 16))
            except ValueError:
                return None
        return None
    return _NAMED_COLORS.get(text.lower())


#: Named paragraph styles -> Word built-in style names.
_NAMED_STYLES = {
    "quote": "Quote",
    "title": "Title",
    "subtitle": "Subtitle",
    "caption": "Caption",
}
#: line-spacing attribute value -> python-docx multiple.
_LINE_SPACING = {"1": 1.0, "1.5": 1.5, "2": 2.0}


def _paragraph_style(paragraph: RichParagraph) -> str | None:
    if paragraph.style == "heading":
        level = min(max(paragraph.level, 1), 6)
        return f"Heading {level}"
    if paragraph.style == "bullet":
        return "List Bullet"
    if paragraph.named_style in _NAMED_STYLES:
        return _NAMED_STYLES[paragraph.named_style]
    return None


def rich_to_docx(document: RichDocument) -> Any:
    """Build a python-docx ``Document`` from a :class:`RichDocument`.

    Raises ``ModuleNotFoundError`` (via the import) when python-docx is absent;
    callers should gate on :func:`python_docx_available` first.
    """
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.shared import Pt, RGBColor

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    out = docx.Document()
    for paragraph in document.paragraphs:
        if paragraph.style == "pagebreak":
            out.add_page_break()
            continue
        para = out.add_paragraph()
        style = _paragraph_style(paragraph)
        if style is not None:
            try:
                para.style = style
            except KeyError:  # pragma: no cover - template without the named style
                pass
        if paragraph.align in align_map:
            para.alignment = align_map[paragraph.align]
        fmt = para.paragraph_format
        if paragraph.line_spacing in _LINE_SPACING:
            fmt.line_spacing = _LINE_SPACING[paragraph.line_spacing]
        if paragraph.space_before:
            fmt.space_before = Pt(paragraph.space_before)
        if paragraph.space_after:
            fmt.space_after = Pt(paragraph.space_after)
        if paragraph.indent:
            fmt.left_indent = Pt(paragraph.indent)
        if paragraph.first_line_indent:
            fmt.first_line_indent = Pt(paragraph.first_line_indent)
        for span in paragraph.spans:
            run = para.add_run(span.text)
            run.bold = span.bold or None
            run.italic = span.italic or None
            run.underline = span.underline or None
            if span.strike:
                run.font.strike = True
            # superscript and subscript share one w:vertAlign element, so set only
            # the active one; assigning the other (even None) would clear it.
            if span.superscript:
                run.font.superscript = True
            elif span.subscript:
                run.font.subscript = True
            if span.font_family:
                run.font.name = span.font_family
            if span.font_size_pt:
                run.font.size = Pt(span.font_size_pt)
            if span.color:
                rgb = _parse_rgb(span.color)
                if rgb is not None:
                    run.font.color.rgb = RGBColor(*rgb)
            if span.highlight:
                name = _HIGHLIGHT_NAMES.get(span.highlight.lower(), "YELLOW")
                run.font.highlight_color = getattr(WD_COLOR_INDEX, name)
    return out


def rich_to_docx_bytes(document: RichDocument) -> bytes:
    """Serialize a :class:`RichDocument` to in-memory ``.docx`` bytes."""
    import io

    buffer = io.BytesIO()
    rich_to_docx(document).save(buffer)
    return buffer.getvalue()


def write_docx(document: Document, target: Path) -> Path:
    """Write ``document``'s canonical markup to ``target`` as a native ``.docx``.

    The markup is converted to the shared :class:`RichDocument` first, so headings,
    bullets, emphasis, and all hidden-codes attributes (font, size, color,
    highlight, alignment) are carried onto Word runs and paragraphs.
    """
    rich = markdown_to_rich(document.text)
    rich_to_docx(rich).save(str(target))
    return target

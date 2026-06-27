"""Native python-docx writer for the hidden-codes formatting feature.

Skipped when python-docx is not installed (it is an optional dependency with a
Pandoc fallback in ``write_docx_document``); when present, the produced ``.docx``
is reopened and its run/paragraph attributes are asserted directly.
"""

from pathlib import Path

import pytest

pytest.importorskip("docx")

from quill.core.document import Document  # noqa: E402
from quill.io.docx_writer import python_docx_available, rich_to_docx, write_docx  # noqa: E402
from quill.io.rtf_model import markdown_to_rich  # noqa: E402


def test_python_docx_available_true_when_installed() -> None:
    assert python_docx_available() is True


def test_run_font_size_color_and_underline() -> None:
    doc = markdown_to_rich('[Hello]{font-family="Arial" font-size="14" color="#C00000" underline}')
    out = rich_to_docx(doc)
    run = out.paragraphs[0].runs[0]
    assert run.text == "Hello"
    assert run.font.name == "Arial"
    assert run.font.size.pt == 14
    assert run.font.color.rgb is not None and str(run.font.color.rgb) == "C00000"
    assert run.underline is True


def test_paragraph_alignment() -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = markdown_to_rich('::: {align="center"}\nCentered.\n:::')
    out = rich_to_docx(doc)
    assert out.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_heading_and_bullet_styles() -> None:
    doc = markdown_to_rich("# Title\n- item")
    out = rich_to_docx(doc)
    assert out.paragraphs[0].style.name == "Heading 1"
    assert out.paragraphs[1].style.name == "List Bullet"


def test_run_strike_super_subscript() -> None:
    out = rich_to_docx(markdown_to_rich("[a]{strike} [b]{superscript} [c]{subscript}"))
    runs = {r.text: r for r in out.paragraphs[0].runs}
    assert runs["a"].font.strike is True
    assert runs["b"].font.superscript is True
    assert runs["c"].font.subscript is True


def test_paragraph_line_spacing_indent_named_style() -> None:
    doc = markdown_to_rich(
        '::: {pstyle="quote" line-spacing="2" indent="36" space-before="12"}\nx\n:::'
    )
    para = rich_to_docx(doc).paragraphs[0]
    assert para.style.name == "Quote"
    assert para.paragraph_format.line_spacing == 2.0
    assert para.paragraph_format.left_indent.pt == 36
    assert para.paragraph_format.space_before.pt == 12


def test_page_break_renders() -> None:
    out = rich_to_docx(markdown_to_rich("a\n::: pagebreak\nb"))
    xml = "".join(p._p.xml for p in out.paragraphs)
    assert "w:br" in xml and 'w:type="page"' in xml


def test_write_docx_roundtrips_through_file(tmp_path: Path) -> None:
    import docx

    target = tmp_path / "out.docx"
    document = Document(text='[Big]{font-size="24"} and plain')
    result = write_docx(document, target)
    assert result == target
    reopened = docx.Document(str(target))
    sizes = [r.font.size.pt for p in reopened.paragraphs for r in p.runs if r.font.size]
    assert 24 in sizes
